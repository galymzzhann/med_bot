# src/embed.py
"""
Section-aware chunking and embedding pipeline.

Key improvements:
  1. Reads structured JSON files (from the new scraper), not raw .txt
  2. Chunks within each section — never mixes "Symptoms" with "Treatment"
  3. Each chunk carries rich metadata: disease name, section type, source file
  4. Uses a multilingual model that actually understands Russian
  5. Deduplicates chunks by content hash

Reads:  data/scraped_json/*.json  +  data/docs/*.{txt,pdf,docx}
Writes: data/embeddings/embeddings.npy  +  data/embeddings/metadata.pkl
"""

import os
import re
import json
import glob
import pickle
import hashlib
import logging
from dataclasses import dataclass, field
from typing import Optional

import yaml
import numpy as np
from sentence_transformers import SentenceTransformer
from pypdf import PdfReader
from docx import Document as DocxDocument


# ── Config ────────────────────────────────────────────────────────────────────

def _load_config() -> dict:
    src  = os.path.dirname(os.path.abspath(__file__))
    root = os.path.dirname(src)
    with open(os.path.join(root, "config.yaml"), encoding="utf-8") as f:
        return yaml.safe_load(f)

def _project_root() -> str:
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


logging.basicConfig(
    format="%(asctime)s %(levelname)s:%(name)s: %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger("embed")

cfg  = _load_config()
ROOT = _project_root()

SCRAPED_DIR = os.path.join(ROOT, cfg["data"]["scraped_dir"])
DOCS_DIR    = os.path.join(ROOT, cfg["data"]["docs_dir"])
EMB_DIR     = os.path.join(ROOT, cfg["data"]["embeddings_dir"])
os.makedirs(EMB_DIR, exist_ok=True)

CHUNK_SIZE  = cfg["embed"]["chunk_size"]
OVERLAP     = cfg["embed"]["chunk_overlap"]
MIN_LEN     = cfg["embed"]["min_chunk_len"]
BATCH_SIZE  = cfg["embed"]["batch_size"]
EMB_MODEL   = cfg["embed"]["model_name"]


# ── Data class ────────────────────────────────────────────────────────────────

@dataclass
class Chunk:
    text:       str
    disease:    str                     # human-readable disease name
    section:    str                     # canonical section: symptoms / diagnostics / …
    source:     str                     # filename
    url:        str                     # original page URL
    chunk_id:   int
    content_hash: str = field(init=False)

    def __post_init__(self):
        self.content_hash = hashlib.md5(self.text.encode("utf-8")).hexdigest()

    def to_meta(self) -> dict:
        return {
            "text":       self.text,
            "disease":    self.disease,
            "section":    self.section,
            "source":     self.source,
            "url":        self.url,
            "chunk_id":   self.chunk_id,
            "hash":       self.content_hash,
        }


# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"[^\w\s\u0400-\u04FF.,;:!?«»()\-–—/%°№@#]", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Chunking (boundary-aware) ─────────────────────────────────────────────────

def _chunk_text(text: str, disease: str, section: str,
                source: str, url: str) -> list[Chunk]:
    """
    Split *text* into overlapping chunks that respect sentence boundaries.
    Every chunk inherits the same disease / section / source metadata.
    """
    text = _clean(text)
    if not text or len(text) < MIN_LEN:
        return []

    # If the entire text fits in one chunk, just return it
    if len(text) <= CHUNK_SIZE:
        return [Chunk(
            text=text, disease=disease, section=section,
            source=source, url=url, chunk_id=0,
        )]

    chunks: list[Chunk] = []
    i   = 0
    cid = 0
    step = CHUNK_SIZE - OVERLAP   # fixed step size (e.g. 800-120 = 680)

    while i < len(text):
        end = min(i + CHUNK_SIZE, len(text))

        # Try to break at a sentence boundary
        if end < len(text):
            boundary = max(
                text.rfind("\n", i, end),
                text.rfind(". ", i, end),
                text.rfind("? ", i, end),
                text.rfind("! ", i, end),
            )
            if boundary > i + CHUNK_SIZE // 3:
                end = boundary + 1

        part = text[i:end].strip()
        if len(part) >= MIN_LEN:
            chunks.append(Chunk(
                text=part, disease=disease, section=section,
                source=source, url=url, chunk_id=cid,
            ))
            cid += 1

        # Always advance by at least `step` to avoid near-duplicates
        i += step

        # If what's left is too small for a standalone chunk, stop
        if len(text) - i < MIN_LEN:
            break

    return chunks


# ── Document loaders ──────────────────────────────────────────────────────────

def _read_txt(path: str) -> str:
    for enc in ("utf-8", "cp1251"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join(
        page.extract_text() or "" for page in reader.pages
    )


def _read_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", c.text).strip()
                     for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ── Load all chunks ──────────────────────────────────────────────────────────

def load_all_chunks() -> list[Chunk]:
    all_chunks:  list[Chunk] = []
    seen_hashes: set[str]    = set()

    # ── 1. Structured JSON files from the new scraper ─────────────────────
    if os.path.isdir(SCRAPED_DIR):
        for path in sorted(glob.glob(os.path.join(SCRAPED_DIR, "*.json"))):
            fname = os.path.basename(path)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    doc = json.load(f)

                disease  = doc.get("title", fname)
                url      = doc.get("url", "")
                sections = doc.get("sections", {})

                file_chunks = 0
                for section_name, section_text in sections.items():
                    if section_name == "_full":
                        # Only use _full if no other sections were extracted
                        continue
                    new = _chunk_text(
                        section_text, disease=disease, section=section_name,
                        source=fname, url=url,
                    )
                    for c in new:
                        if c.content_hash not in seen_hashes:
                            seen_hashes.add(c.content_hash)
                            all_chunks.append(c)
                            file_chunks += 1

                # Fallback: if no named sections, chunk _full as "general"
                if file_chunks == 0 and "_full" in sections:
                    new = _chunk_text(
                        sections["_full"], disease=disease, section="general",
                        source=fname, url=url,
                    )
                    for c in new:
                        if c.content_hash not in seen_hashes:
                            seen_hashes.add(c.content_hash)
                            all_chunks.append(c)
                            file_chunks += 1

                logger.info(f"JSON  {fname}: {file_chunks} chunks")

            except Exception:
                logger.exception(f"Failed to read JSON: {fname}")

    # ── 2. Legacy / additional docs (txt, pdf, docx) ─────────────────────
    if os.path.isdir(DOCS_DIR):
        readers = [("*.txt", _read_txt), ("*.pdf", _read_pdf), ("*.docx", _read_docx)]
        for pattern, reader_fn in readers:
            for path in sorted(glob.glob(os.path.join(DOCS_DIR, pattern))):
                fname = os.path.basename(path)
                try:
                    raw = reader_fn(path).strip()
                    if not raw:
                        continue
                    new = _chunk_text(
                        raw, disease=fname, section="general",
                        source=fname, url="",
                    )
                    deduped = [c for c in new if c.content_hash not in seen_hashes]
                    for c in deduped:
                        seen_hashes.add(c.content_hash)
                    all_chunks.extend(deduped)
                    logger.info(f"DOC   {fname}: {len(deduped)} chunks")
                except Exception:
                    logger.exception(f"Failed to read doc: {fname}")

    logger.info(f"Total chunks: {len(all_chunks)}")
    return all_chunks


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    chunks = load_all_chunks()
    if not chunks:
        raise RuntimeError(
            f"No chunks found. Ensure files exist in '{SCRAPED_DIR}' or '{DOCS_DIR}'."
        )

    # Section labels for enriched embedding text
    SECTION_LABELS = {
        "definition": "Определение",
        "symptoms": "Симптомы",
        "diagnostics": "Диагностика",
        "treatment": "Лечение",
        "prevention": "Профилактика",
        "classification": "Классификация",
        "etiology": "Этиология",
        "general": "Общая информация",
    }

    # Contextual enrichment: prepend disease name + section to each chunk
    # so the embedding captures WHICH disease the text belongs to.
    # Without this, "боль, рвота, температура" matches any abdominal disease.
    # With this, "Острый панкреатит. Симптомы: боль, рвота, температура"
    # matches pancreatitis specifically.
    texts = [
        f"{c.disease}. {SECTION_LABELS.get(c.section, c.section)}: {c.text}"
        for c in chunks
    ]
    total = len(texts)

    logger.info(f"Loading multilingual embedding model: {EMB_MODEL}")
    model = SentenceTransformer(EMB_MODEL)

    logger.info(f"Encoding {total} chunks (batch_size={BATCH_SIZE}) …")
    all_embs: list[np.ndarray] = []
    for start in range(0, total, BATCH_SIZE):
        batch = texts[start : start + BATCH_SIZE]
        emb   = model.encode(
            batch, convert_to_numpy=True,
            normalize_embeddings=True, show_progress_bar=False,
        )
        all_embs.append(emb)
        done = min(start + BATCH_SIZE, total)
        logger.info(f"  {done}/{total} ({int(done / total * 100)}%)")

    embeddings = np.vstack(all_embs).astype("float32")

    emb_path  = os.path.join(EMB_DIR, "embeddings.npy")
    meta_path = os.path.join(EMB_DIR, "metadata.pkl")

    np.save(emb_path, embeddings)
    meta = [c.to_meta() for c in chunks]
    with open(meta_path, "wb") as f:
        pickle.dump(meta, f)

    logger.info(f"Saved embeddings {embeddings.shape} → {emb_path}")
    logger.info(f"Saved metadata ({len(meta)} entries) → {meta_path}")


if __name__ == "__main__":
    main()